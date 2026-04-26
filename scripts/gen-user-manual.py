"""Generate USER-MANUAL.pdf and USER-MANUAL.docx from USER-MANUAL.md.
Minimal converter: reads markdown, emits a readable PDF + DOCX with headings
and paragraphs. Not a full-fidelity renderer; meant as a baseline so the
artifacts exist."""
import re
from pathlib import Path

from docx import Document  # type: ignore
from docx.shared import Inches  # type: ignore
from reportlab.lib.pagesizes import LETTER  # type: ignore
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
from reportlab.lib.units import inch  # type: ignore
from reportlab.platypus import (  # type: ignore
    Image as RLImage,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
)

ROOT = Path(__file__).resolve().parent.parent
MD = (ROOT / "USER-MANUAL.md").read_text(encoding="utf-8")

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def parse(md):
    blocks = []
    in_code = False
    buf = []
    for line in md.splitlines():
        if line.startswith("```"):
            if in_code:
                blocks.append(("code", "\n".join(buf)))
                buf = []
                in_code = False
            else:
                if buf:
                    blocks.append(("p", "\n".join(buf).strip()))
                    buf = []
                in_code = True
            continue
        if in_code:
            buf.append(line)
            continue
        m = IMAGE_RE.match(line)
        if m:
            if buf:
                blocks.append(("p", "\n".join(buf).strip()))
                buf = []
            blocks.append(("image", f"{m.group(1)}|{m.group(2)}"))
            continue
        if line.startswith("# "):
            if buf:
                blocks.append(("p", "\n".join(buf).strip()))
                buf = []
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            if buf:
                blocks.append(("p", "\n".join(buf).strip()))
                buf = []
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            if buf:
                blocks.append(("p", "\n".join(buf).strip()))
                buf = []
            blocks.append(("h3", line[4:].strip()))
        elif line.strip() == "":
            if buf:
                blocks.append(("p", "\n".join(buf).strip()))
                buf = []
        else:
            buf.append(line)
    if buf:
        blocks.append(("p", "\n".join(buf).strip()))
    return [b for b in blocks if b[1]]


def _resolve_image(path: str) -> Path:
    """Resolve image path relative to ROOT. Prefer .png over .svg."""
    p = ROOT / path
    if p.suffix.lower() == ".svg":
        png = p.with_suffix(".png")
        if png.exists():
            return png
    return p


blocks = parse(MD)

# DOCX
doc = Document()
for kind, text in blocks:
    if kind == "h1":
        doc.add_heading(text, level=1)
    elif kind == "h2":
        doc.add_heading(text, level=2)
    elif kind == "h3":
        doc.add_heading(text, level=3)
    elif kind == "code":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
    elif kind == "image":
        alt, path = text.split("|", 1)
        img_path = _resolve_image(path)
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(6))
            except Exception as e:
                print(f"WARN: docx image embed failed for {img_path}: {e}")
                doc.add_paragraph(f"[image: {alt} ({path})]")
        else:
            print(f"WARN: image not found: {img_path}")
            doc.add_paragraph(f"[image: {alt} ({path})]")
    else:
        doc.add_paragraph(text)
doc.save(ROOT / "USER-MANUAL.docx")

# PDF
styles = getSampleStyleSheet()
code_style = ParagraphStyle("code", parent=styles["Code"], fontSize=8, leading=10)
pdf = SimpleDocTemplate(
    str(ROOT / "USER-MANUAL.pdf"),
    pagesize=LETTER,
    leftMargin=0.9 * inch,
    rightMargin=0.9 * inch,
    topMargin=0.9 * inch,
    bottomMargin=0.9 * inch,
    title="CivicSuite User Manual",
)
flow = []
for kind, text in blocks:
    safe = (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))
    if kind == "h1":
        flow.append(Paragraph(safe, styles["Title"]))
    elif kind == "h2":
        flow.append(Paragraph(safe, styles["Heading1"]))
    elif kind == "h3":
        flow.append(Paragraph(safe, styles["Heading2"]))
    elif kind == "code":
        flow.append(Preformatted(text, code_style))
    elif kind == "image":
        alt, path = text.split("|", 1)
        img_path = _resolve_image(path)
        if img_path.exists():
            try:
                from PIL import Image as PILImage  # type: ignore
                with PILImage.open(str(img_path)) as im:
                    iw, ih = im.size
                target_w = 6.0 * inch
                target_h = target_w * (ih / iw)
                flow.append(RLImage(str(img_path), width=target_w, height=target_h))
            except Exception as e:
                print(f"WARN: pdf image embed failed for {img_path}: {e}")
                flow.append(Paragraph(f"[image: {alt} ({path})]", styles["BodyText"]))
        else:
            print(f"WARN: image not found: {img_path}")
            flow.append(Paragraph(f"[image: {alt} ({path})]", styles["BodyText"]))
    else:
        flow.append(Paragraph(safe, styles["BodyText"]))
    flow.append(Spacer(1, 6))
pdf.build(flow)
print("OK: wrote USER-MANUAL.pdf and USER-MANUAL.docx")
